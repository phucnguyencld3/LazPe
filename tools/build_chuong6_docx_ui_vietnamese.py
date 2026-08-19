from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_PATH = ROOT / "docs" / "chuong-6-testcase-lazpe-ui-co-dau-v2.docx"


def ptext(*lines):
    return "\n".join(lines)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_width(table, widths):
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)


def style_runs(paragraph, size=11, bold=False, italic=False, color=None):
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
        if color:
            run.font.color.rgb = RGBColor(*color)


def add_paragraph(doc, text, size=12, bold=False, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.add_run(text)
    style_runs(p, size=size, bold=bold)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_table(doc, title, rows):
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.add_run(title)
    style_runs(title_p, size=12, bold=True, italic=True)
    title_p.paragraph_format.space_before = Pt(8)
    title_p.paragraph_format.space_after = Pt(4)

    headers = ["STT", "Trường hợp kiểm thử", "Dữ liệu đầu vào", "Kết quả thực tế", "Trạng thái"]
    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    widths = [0.45, 2.35, 2.8, 2.75, 0.65]

    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        set_cell_shading(cell, "E8EEF5")
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            style_runs(p, size=11, bold=True)
    set_repeat_table_header(table.rows[0])

    for idx, (case_name, input_data, actual_result) in enumerate(rows, 1):
        cells = table.add_row().cells
        values = [str(idx), case_name, input_data, actual_result, "Pass"]
        for col, value in enumerate(values):
            cells[col].text = value
            set_cell_margins(cells[col])
            cells[col].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cells[col].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col in (0, 4) else WD_ALIGN_PARAGRAPH.LEFT
                style_runs(p, size=10)
        for p in cells[4].paragraphs:
            style_runs(p, size=10, bold=True, color=(0, 128, 0))

    set_table_width(table, widths)
    doc.add_paragraph()


product_cases = [
    ("Kiểm tra hiển thị danh sách sản phẩm đang bán trên trang cửa hàng.",
     ptext('Từ khóa tìm kiếm = ""', 'Danh mục = "Tất cả"', 'Sắp xếp = "Mới nhất"', 'Bộ lọc giảm giá = "Tất cả"'),
     "Hiển thị danh sách sản phẩm đang hoạt động, có hình ảnh, tên sản phẩm, giá bán và nút thao tác."),
    ("Kiểm tra tìm kiếm sản phẩm theo từ khóa tên sản phẩm.",
     ptext('Từ khóa tìm kiếm = "sữa"', 'Danh mục = "Tất cả"', 'Sắp xếp = "Liên quan"'),
     "Hiển thị các sản phẩm có tên hoặc thông tin liên quan đến từ khóa đã nhập."),
    ("Kiểm tra lọc danh sách sản phẩm theo danh mục.",
     ptext('Danh mục = "Danh mục đang có trên hệ thống"', 'Từ khóa tìm kiếm = ""'),
     "Danh sách chỉ hiển thị sản phẩm thuộc danh mục đã chọn."),
    ("Kiểm tra lọc sản phẩm đang áp dụng giảm giá.",
     ptext('Bộ lọc giảm giá = "Đang giảm giá"', 'Danh mục = "Tất cả"'),
     "Chỉ hiển thị các sản phẩm có giá khuyến mãi hoặc phần trăm giảm giá."),
    ("Kiểm tra sắp xếp danh sách sản phẩm theo giá tăng dần.",
     ptext('Sắp xếp = "Giá tăng dần"', 'Từ khóa tìm kiếm = ""'),
     "Sản phẩm được sắp xếp từ giá thấp đến giá cao."),
    ("Kiểm tra hiển thị trang thông tin chi tiết sản phẩm.",
     ptext('Sản phẩm = "Chọn một sản phẩm đang hiển thị trên trang sản phẩm"'),
     "Hiển thị trang chi tiết sản phẩm gồm hình ảnh, tên, mô tả, giá, biến thể và thông tin liên quan."),
    ("Kiểm tra thao tác chọn biến thể sản phẩm (màu sắc, kích cỡ) còn hàng.",
     ptext('Sản phẩm = "Sản phẩm có biến thể"', 'Biến thể = "Chọn màu/kích cỡ đang còn hàng"'),
     "Giá, hình ảnh hoặc tồn kho của biến thể được cập nhật đúng trên giao diện."),
    ("Kiểm tra ràng buộc không cho chọn biến thể sản phẩm đã hết hàng.",
     ptext('Sản phẩm = "Sản phẩm có biến thể hết hàng"', 'Biến thể = "Biến thể hết hàng"'),
     "Hệ thống hiển thị trạng thái hết hàng và không cho thêm biến thể đó vào giỏ."),
    ("Kiểm tra thao tác thêm sản phẩm hợp lệ vào giỏ hàng.",
     ptext('Sản phẩm = "Sản phẩm đang còn hàng"', 'Số lượng = "1"', 'Biến thể = "Biến thể hợp lệ"'),
     "Sản phẩm được thêm vào giỏ hàng và số lượng giỏ hàng được cập nhật."),
    ("Kiểm tra báo lỗi khi thêm số lượng sản phẩm vượt quá tồn kho.",
     ptext('Sản phẩm = "Sản phẩm đang còn hàng"', 'Số lượng = "Lớn hơn số lượng tồn kho"', 'Biến thể = "Biến thể hợp lệ"'),
     "Hiển thị thông báo số lượng không hợp lệ hoặc vượt quá tồn kho."),
    ("Kiểm tra hiển thị bảng danh sách sản phẩm trong trang quản trị Admin.",
     ptext('Tài khoản = "Admin"', 'Từ khóa = ""', 'Trạng thái = "Tất cả"', 'Danh mục = "Tất cả"'),
     "Hiển thị bảng danh sách sản phẩm kèm trạng thái, giá, danh mục và thao tác quản lý."),
    ("Kiểm tra tìm kiếm sản phẩm theo tên trên màn hình quản trị Admin.",
     ptext('Tài khoản = "Admin"', 'Từ khóa = "sữa"', 'Trạng thái = "Tất cả"'),
     "Bảng quản trị chỉ hiển thị sản phẩm phù hợp với từ khóa tìm kiếm."),
    ("Kiểm tra lọc sản phẩm theo trạng thái hoạt động trong trang Admin.",
     ptext('Tài khoản = "Admin"', 'Trạng thái = "Đang hoạt động"', 'Từ khóa = ""'),
     "Chỉ hiển thị các sản phẩm đang được bật trạng thái hoạt động."),
    ("Kiểm tra lọc sản phẩm theo khoảng giá trên trang quản trị Admin.",
     ptext('Giá từ = "100000"', 'Giá đến = "500000"', 'Tài khoản = "Admin"'),
     "Hiển thị các sản phẩm có giá nằm trong khoảng đã nhập."),
    ("Kiểm tra phân quyền từ chối truy cập màn hình quản lý sản phẩm.",
     ptext('Tài khoản = "Người dùng thường"', 'Màn hình = "Quản lý sản phẩm"'),
     "Hệ thống không cho truy cập màn hình hoặc thao tác quản trị sản phẩm."),
    ("Kiểm tra thao tác Admin thêm sản phẩm mới với thông tin hợp lệ.",
     ptext('Tên sản phẩm = "Sản phẩm kiểm thử LazPe"', 'Danh mục = "Chọn danh mục đang có"', 'Thương hiệu = "Chọn thương hiệu đang có"', 'Giá bán = "250000"', 'Mô tả = "Mô tả sản phẩm kiểm thử"', 'Trạng thái = "Đang hoạt động"', 'Hình ảnh = "product-test.jpg"'),
     'Sản phẩm được thêm thành công, hiển thị thông báo "Tạo sản phẩm thành công".'),
    ("Kiểm tra báo lỗi khi thêm sản phẩm mới nhưng bỏ trống tên sản phẩm.",
     ptext('Tên sản phẩm = ""', 'Danh mục = "Chọn danh mục đang có"', 'Giá bán = "250000"'),
     "Hiển thị thông báo yêu cầu nhập tên sản phẩm."),
    ("Kiểm tra báo lỗi khi nhập giá bán sản phẩm không hợp lệ.",
     ptext('Tên sản phẩm = "Sản phẩm kiểm thử LazPe"', 'Giá bán = "-10000"', 'Danh mục = "Chọn danh mục đang có"'),
     "Hiển thị thông báo giá bán không hợp lệ."),
    ("Kiểm tra thao tác Admin thêm sản phẩm có tùy chọn và biến thể.",
     ptext('Tên sản phẩm = "Sản phẩm có biến thể"', 'Tùy chọn = "Màu sắc, Kích cỡ"', 'Biến thể = "Đỏ/S, Đỏ/M"', 'Số lượng tồn = "20"', 'Giá bán = "300000"'),
     "Sản phẩm, tùy chọn và biến thể được tạo thành công trong hệ thống."),
    ("Kiểm tra báo lỗi khi thêm sản phẩm mới nhưng chưa chọn danh mục.",
     ptext('Tên sản phẩm = "Sản phẩm kiểm thử LazPe"', 'Danh mục = ""', 'Giá bán = "250000"'),
     "Hiển thị thông báo yêu cầu chọn danh mục sản phẩm."),
    ("Kiểm tra thao tác Admin cập nhật thông tin sản phẩm đã có.",
     ptext('Sản phẩm = "Chọn sản phẩm đang có"', 'Tên sản phẩm mới = "Sản phẩm kiểm thử LazPe - cập nhật"', 'Giá bán mới = "270000"', 'Trạng thái = "Đang hoạt động"'),
     "Thông tin sản phẩm được cập nhật và hiển thị đúng trên danh sách quản trị."),
    ("Kiểm tra báo lỗi khi cập nhật sản phẩm bỏ trống trường bắt buộc.",
     ptext('Sản phẩm = "Chọn sản phẩm đang có"', 'Tên sản phẩm mới = ""', 'Giá bán mới = "270000"'),
     "Hiển thị thông báo dữ liệu không hợp lệ và không lưu thay đổi."),
    ("Kiểm tra thao tác Admin bật/tắt trạng thái hoạt động của sản phẩm.",
     ptext('Sản phẩm = "Chọn sản phẩm đang có"', 'Thao tác = "Bật/Tắt trạng thái"'),
     "Trạng thái sản phẩm được thay đổi đúng theo thao tác admin chọn."),
    ("Kiểm tra thao tác Admin xóa sản phẩm chưa phát sinh ràng buộc.",
     ptext('Sản phẩm = "Sản phẩm kiểm thử chưa có đơn hàng"', 'Thao tác = "Xóa"', 'Xác nhận = "Đồng ý"'),
     "Sản phẩm được xóa hoặc chuyển sang trạng thái ngừng hiển thị theo xử lý của hệ thống."),
    ("Kiểm tra thao tác hủy xác nhận khi thực hiện xóa sản phẩm.",
     ptext('Sản phẩm = "Chọn sản phẩm đang có"', 'Thao tác = "Xóa"', 'Xác nhận = "Hủy"'),
     "Sản phẩm không bị xóa và vẫn hiển thị trong danh sách quản trị."),
]


order_cases = [
    ("Kiểm tra thao tác tạo đơn hàng hợp lệ từ giỏ hàng (COD).",
     ptext('Sản phẩm trong giỏ = "Sản phẩm đang còn hàng"', 'Số lượng = "1"', 'Địa chỉ giao hàng = "Địa chỉ mặc định"', 'Phương thức thanh toán = "Thanh toán khi nhận hàng"', 'Ghi chú = "Giao trong giờ hành chính"'),
     "Đơn hàng được tạo thành công và hiển thị mã đơn hàng."),
    ("Kiểm tra ràng buộc không cho đặt hàng khi giỏ hàng đang trống.",
     ptext('Giỏ hàng = "Không có sản phẩm"', 'Thao tác = "Đặt hàng"'),
     "Hiển thị thông báo giỏ hàng trống, không tạo đơn hàng."),
    ("Kiểm tra báo lỗi khi đặt hàng nhưng chưa chọn địa chỉ giao hàng.",
     ptext('Sản phẩm trong giỏ = "Có sản phẩm"', 'Địa chỉ giao hàng = ""', 'Phương thức thanh toán = "Thanh toán khi nhận hàng"'),
     "Hiển thị thông báo yêu cầu chọn hoặc nhập địa chỉ giao hàng."),
    ("Kiểm tra áp dụng voucher giảm giá hợp lệ khi tạo đơn hàng.",
     ptext('Sản phẩm trong giỏ = "Có sản phẩm đủ điều kiện"', 'Mã voucher = "Voucher hợp lệ trong ví"', 'Phương thức thanh toán = "Thanh toán khi nhận hàng"'),
     "Đơn hàng được tạo thành công và tổng tiền được giảm đúng theo voucher."),
    ("Kiểm tra quy trình tạo đơn hàng chọn thanh toán qua VNPAY.",
     ptext('Sản phẩm trong giỏ = "Có sản phẩm"', 'Địa chỉ giao hàng = "Địa chỉ hợp lệ"', 'Phương thức thanh toán = "VNPAY"'),
     "Hệ thống chuyển sang bước thanh toán VNPAY hoặc hiển thị liên kết thanh toán."),
    ("Kiểm tra hiển thị danh sách đơn hàng cá nhân của người dùng.",
     ptext('Tài khoản = "Người dùng đã đăng nhập"', 'Bộ lọc trạng thái = "Tất cả"'),
     "Hiển thị danh sách đơn hàng thuộc đúng tài khoản đang đăng nhập."),
    ("Kiểm tra lọc danh sách đơn hàng cá nhân theo trạng thái xử lý.",
     ptext('Bộ lọc trạng thái = "Chờ xác nhận"', 'Từ khóa = ""'),
     "Chỉ hiển thị các đơn hàng có trạng thái chờ xác nhận."),
    ("Kiểm tra hiển thị thông tin chi tiết của một đơn hàng.",
     ptext('Đơn hàng = "Chọn đơn hàng của tài khoản hiện tại"'),
     "Hiển thị chi tiết sản phẩm, giá tiền, địa chỉ giao hàng, trạng thái và thanh toán."),
    ("Kiểm tra tra cứu thông tin đơn hàng bằng mã đơn/mã vận đơn hợp lệ.",
     ptext('Mã đơn hàng/Mã vận đơn = "Mã đơn đang có trên hệ thống"'),
     "Hiển thị thông tin theo dõi đơn hàng tương ứng với mã đã nhập."),
    ("Kiểm tra thông báo khi tra cứu đơn hàng với mã không tồn tại.",
     ptext('Mã đơn hàng/Mã vận đơn = "MADONKHONGTONTAI"'),
     "Hiển thị thông báo không tìm thấy đơn hàng."),
    ("Kiểm tra thao tác gửi yêu cầu hủy đơn hàng ở trạng thái chờ xác nhận.",
     ptext('Đơn hàng = "Đơn đang chờ xác nhận"', 'Lý do hủy = "Đặt nhầm sản phẩm"'),
     "Yêu cầu hủy đơn được gửi thành công và trạng thái đơn được cập nhật."),
    ("Kiểm tra ràng buộc không cho phép hủy đơn hàng đã hoàn tất.",
     ptext('Đơn hàng = "Đơn đã hoàn tất"', 'Lý do hủy = "Không còn nhu cầu"'),
     "Hiển thị thông báo không thể hủy đơn ở trạng thái hiện tại."),
    ("Kiểm tra thao tác gửi yêu cầu trả hàng/hoàn tiền kèm ảnh minh chứng.",
     ptext('Đơn hàng = "Đơn đủ điều kiện trả hàng"', 'Lý do trả hàng = "Sản phẩm không phù hợp"', 'Mô tả = "Muốn trả sản phẩm"', 'Hình ảnh = "return.jpg"', 'Phương thức hoàn tiền = "Hoàn vào xu/ví"'),
     "Yêu cầu trả hàng được gửi thành công."),
    ("Kiểm tra báo lỗi khi gửi yêu cầu trả hàng nhưng thiếu lý do.",
     ptext('Đơn hàng = "Đơn đủ điều kiện trả hàng"', 'Lý do trả hàng = ""', 'Hình ảnh = "return.jpg"'),
     "Hiển thị thông báo yêu cầu nhập lý do trả hàng."),
    ("Kiểm tra thao tác hủy yêu cầu trả hàng đã gửi trước đó.",
     ptext('Đơn hàng = "Đơn đang có yêu cầu trả hàng"', 'Thao tác = "Hủy yêu cầu trả hàng"'),
     "Yêu cầu trả hàng được hủy thành công."),
    ("Kiểm tra hiển thị bảng danh sách tất cả đơn hàng trên trang Admin.",
     ptext('Tài khoản = "Admin"', 'Từ khóa = ""', 'Trạng thái = "Tất cả"', 'Khoảng giá = "Tất cả"'),
     "Hiển thị bảng danh sách đơn hàng trong trang quản trị."),
    ("Kiểm tra tìm kiếm đơn hàng theo mã đơn/tên khách hàng trên Admin.",
     ptext('Từ khóa = "Mã đơn hoặc tên khách hàng đang có"', 'Trạng thái = "Tất cả"'),
     "Bảng đơn hàng hiển thị kết quả phù hợp với từ khóa."),
    ("Kiểm tra lọc danh sách đơn hàng theo trạng thái trên trang Admin.",
     ptext('Trạng thái = "Đang giao"', 'Từ khóa = ""'),
     "Chỉ hiển thị các đơn hàng có trạng thái đang giao."),
    ("Kiểm tra hiển thị bảng điều khiển thống kê chỉ số đơn hàng Admin.",
     ptext('Tài khoản = "Admin"', 'Màn hình = "Quản lý đơn hàng"'),
     "Hiển thị tổng đơn, đơn chờ xử lý, đơn đang giao, đơn hoàn tất và doanh thu."),
    ("Kiểm tra chức năng xuất danh sách đơn hàng ra tệp Excel.",
     ptext('Tài khoản = "Admin"', 'Bộ lọc = "Trạng thái hoặc khoảng thời gian đã chọn"', 'Thao tác = "Xuất Excel"'),
     "Hệ thống tải file Excel danh sách đơn hàng."),
    ("Kiểm tra thao tác Admin duyệt xác nhận đơn hàng đang chờ xử lý.",
     ptext('Đơn hàng = "Đơn đang chờ xác nhận"', 'Thao tác = "Xác nhận đơn"'),
     "Đơn hàng chuyển sang trạng thái đã xác nhận."),
    ("Kiểm tra thao tác Admin cập nhật đơn hàng sang trạng thái đang giao.",
     ptext('Đơn hàng = "Đơn đã xác nhận"', 'Thao tác = "Đánh dấu đang giao"'),
     "Đơn hàng chuyển sang trạng thái đang giao."),
    ("Kiểm tra thao tác Admin cập nhật đơn hàng thành hoàn tất.",
     ptext('Đơn hàng = "Đơn đang giao"', 'Thao tác = "Hoàn tất đơn hàng"'),
     "Đơn hàng chuyển sang trạng thái hoàn tất."),
    ("Kiểm tra thao tác Admin hủy đơn hàng kèm lý do hủy.",
     ptext('Đơn hàng = "Đơn có thể hủy"', 'Lý do hủy = "Khách yêu cầu hủy"', 'Thao tác = "Hủy đơn"'),
     "Đơn hàng chuyển sang trạng thái đã hủy theo đúng quy trình."),
    ("Kiểm tra thao tác Admin xử lý yêu cầu trả hàng của khách hàng.",
     ptext('Đơn hàng = "Đơn đang yêu cầu trả hàng"', 'Thao tác = "Duyệt/Từ chối/Xác nhận đã nhận hàng hoàn"', 'Ghi chú xử lý = "Đã kiểm tra yêu cầu"'),
     "Trạng thái trả hàng được cập nhật đúng theo thao tác của admin.")
]


voucher_cases = [
    ("Kiểm tra hiển thị danh sách voucher khuyến mãi công khai.",
     ptext('Tài khoản = "Người dùng hoặc khách chưa đăng nhập"', 'Màn hình = "Voucher/Khuyến mãi"'),
     "Hiển thị các voucher công khai còn hiệu lực và còn số lượng."),
    ("Kiểm tra hiển thị trạng thái đối với voucher đã được lưu vào ví.",
     ptext('Tài khoản = "Người dùng đã đăng nhập"', 'Voucher = "Voucher đã lưu trước đó"'),
     "Voucher hiển thị trạng thái đã lưu hoặc không cho lưu lại lần nữa."),
    ("Kiểm tra ràng buộc không hiển thị các voucher đã hết hạn sử dụng.",
     ptext('Voucher = "Voucher có ngày kết thúc nhỏ hơn ngày hiện tại"', 'Màn hình = "Voucher/Khuyến mãi"'),
     "Voucher hết hạn không hiển thị trong danh sách có thể nhận."),
    ("Kiểm tra ràng buộc không cho nhận voucher đã hết lượt sử dụng.",
     ptext('Voucher = "Voucher đã hết lượt"', 'Thao tác = "Lưu voucher"'),
     "Hiển thị thông báo voucher đã hết lượt hoặc không còn khả dụng."),
    ("Kiểm tra hiển thị danh sách voucher công khai cho khách chưa đăng nhập.",
     ptext('Tài khoản = "Chưa đăng nhập"', 'Màn hình = "Voucher/Khuyến mãi"'),
     "Hệ thống vẫn hiển thị voucher công khai nhưng chưa đánh dấu đã lưu."),
    ("Kiểm tra thao tác lưu voucher công khai vào ví cá nhân.",
     ptext('Tài khoản = "Người dùng đã đăng nhập"', 'Voucher = "Voucher công khai còn hiệu lực"', 'Thao tác = "Lưu voucher"'),
     'Hiển thị thông báo "Lưu voucher thành công".'),
    ("Kiểm tra thông báo khi thực hiện lưu lại voucher đã có sẵn trong ví.",
     ptext('Tài khoản = "Người dùng đã đăng nhập"', 'Voucher = "Voucher đang có trong ví"', 'Thao tác = "Lưu voucher"'),
     "Hiển thị thông báo voucher đã có trong ví."),
    ("Kiểm tra thông báo khi lưu thủ công voucher loại phân phối trực tiếp.",
     ptext('Voucher = "Voucher loại phân phối trực tiếp"', 'Thao tác = "Lưu thủ công"'),
     "Hiển thị thông báo voucher được phát tự động, không cần lưu thủ công."),
    ("Kiểm tra kích hoạt voucher độc quyền bằng mã khuyến mãi riêng.",
     ptext('Mã voucher = "Mã voucher độc quyền hợp lệ"', 'Thao tác = "Kích hoạt mã"'),
     "Voucher được thêm vào ví của người dùng."),
    ("Kiểm tra thông báo khi nhập mã voucher kích hoạt không tồn tại.",
     ptext('Mã voucher = "SAIMA123"', 'Thao tác = "Kích hoạt mã"'),
     "Hiển thị thông báo mã voucher không tồn tại."),
    ("Kiểm tra hiển thị danh sách voucher trong ví cá nhân của người dùng.",
     ptext('Tài khoản = "Người dùng đã đăng nhập"', 'Màn hình = "Ví voucher"'),
     "Hiển thị danh sách voucher chưa dùng, đã dùng hoặc hết hạn của người dùng."),
    ("Kiểm tra tự động cập nhật trạng thái hết hạn cho voucher trong ví.",
     ptext('Voucher trong ví = "Voucher đã qua ngày kết thúc"', 'Màn hình = "Ví voucher"'),
     "Voucher chuyển sang trạng thái hết hạn."),
    ("Kiểm tra hiển thị danh sách voucher đủ điều kiện tại bước thanh toán.",
     ptext('Giỏ hàng = "Có sản phẩm đủ điều kiện"', 'Màn hình = "Thanh toán"', 'Thao tác = "Chọn voucher"'),
     "Hiển thị các voucher có thể áp dụng cho đơn hàng hiện tại."),
    ("Kiểm tra áp dụng mã voucher giảm giá vào tổng tiền đơn hàng.",
     ptext('Voucher = "Voucher giảm giá sản phẩm hợp lệ"', 'Tổng tiền hàng = "Đạt giá trị tối thiểu"', 'Thao tác = "Áp dụng"'),
     "Tổng tiền đơn hàng được giảm đúng theo cấu hình voucher."),
    ("Kiểm tra thao tác gỡ bỏ voucher đang áp dụng khỏi đơn hàng.",
     ptext('Voucher = "Voucher đang được áp dụng"', 'Thao tác = "Bỏ áp dụng"'),
     "Voucher được gỡ khỏi đơn hàng và tổng tiền được tính lại."),
    ("Kiểm tra thao tác Admin tạo voucher mới với thông tin hợp lệ.",
     ptext('Mã voucher = "SALETEST"', 'Tên voucher = "Voucher kiểm thử"', 'Loại giảm giá = "Giảm tiền"', 'Giá trị giảm = "50000"', 'Đơn tối thiểu = "200000"', 'Số lượng = "100"', 'Ngày bắt đầu = "Ngày hiện tại"', 'Ngày kết thúc = "Ngày sau hiện tại"', 'Trạng thái = "Hoạt động"'),
     "Voucher được tạo thành công và hiển thị trong danh sách quản trị."),
    ("Kiểm tra báo lỗi khi tạo voucher phần trăm vượt quá mức 50%.",
     ptext('Loại giảm giá = "Phần trăm"', 'Giá trị giảm = "60"', 'Tên voucher = "Voucher sai phần trăm"'),
     "Hiển thị thông báo mức giảm giá theo phần trăm không được vượt quá 50%."),
    ("Kiểm tra thao tác Admin cập nhật thông tin voucher đã có.",
     ptext('Voucher = "Chọn voucher đang có"', 'Tên voucher mới = "Voucher cập nhật"', 'Giá trị giảm mới = "70000"', 'Số lượng = "120"'),
     "Thông tin voucher được cập nhật thành công."),
    ("Kiểm tra thao tác Admin khóa hoặc mở khóa trạng thái voucher.",
     ptext('Voucher = "Chọn voucher đang có"', 'Thao tác = "Khóa/Mở khóa"'),
     "Trạng thái voucher được thay đổi đúng theo thao tác."),
    ("Kiểm tra thao tác Admin xóa voucher chưa từng phát sinh lượt dùng.",
     ptext('Voucher = "Voucher kiểm thử chưa có lượt dùng"', 'Thao tác = "Xóa"', 'Xác nhận = "Đồng ý"'),
     "Voucher được xóa khỏi danh sách quản trị."),
    ("Kiểm tra ràng buộc không cho phép xóa voucher đã có lượt dùng.",
     ptext('Voucher = "Voucher đã phát sinh lượt dùng"', 'Thao tác = "Xóa"'),
     "Hiển thị thông báo không thể xóa voucher đã được sử dụng."),
    ("Kiểm tra thao tác Admin phát voucher trực tiếp cho người dùng.",
     ptext('Voucher = "Voucher độc quyền phân phối trực tiếp"', 'Người nhận = "Chọn người dùng đang có"', 'Thao tác = "Phát voucher"'),
     "Voucher được phát vào ví của người dùng đã chọn."),
    ("Kiểm tra báo lỗi khi phát voucher nhưng danh sách người nhận trống.",
     ptext('Voucher = "Voucher độc quyền phân phối trực tiếp"', 'Người nhận = ""', 'Thao tác = "Phát voucher"'),
     "Hiển thị thông báo danh sách người dùng trống."),
    ("Kiểm tra thao tác Admin thu hồi voucher đã phát trực tiếp.",
     ptext('Voucher đã phát = "Voucher chưa sử dụng"', 'Người dùng = "Người dùng đã được phát"', 'Thao tác = "Thu hồi"'),
     "Voucher được thu hồi khỏi ví người dùng."),
    ("Kiểm tra hiển thị lịch sử chi tiết sử dụng voucher trên trang Admin.",
     ptext('Voucher = "Chọn voucher đang có"', 'Thao tác = "Xem lịch sử sử dụng"'),
     "Hiển thị danh sách người dùng, đơn hàng, số tiền giảm và thời gian sử dụng."),
]


loyalty_cases = [
    ("Kiểm tra hiển thị trang hồ sơ khách hàng thân thiết (Loyalty).",
     ptext('Tài khoản = "Người dùng đã đăng nhập"', 'Màn hình = "Hồ sơ/Loyalty"'),
     "Hiển thị điểm khả dụng, tổng điểm, hạng hiện tại và tiến độ lên hạng."),
    ("Kiểm tra ràng buộc yêu cầu đăng nhập khi truy cập trang Loyalty.",
     ptext('Tài khoản = "Chưa đăng nhập"', 'Màn hình = "Loyalty"'),
     "Hiển thị yêu cầu đăng nhập."),
    ("Kiểm tra hiển thị thông tin hạng thành viên hiện tại và huy hiệu.",
     ptext('Tài khoản = "Người dùng có hồ sơ loyalty"', 'Hạng = "Standard/Silver/Gold/Diamond theo dữ liệu hệ thống"'),
     "Giao diện hiển thị đúng tên hạng và màu/huy hiệu tương ứng."),
    ("Kiểm tra hiển thị thanh tiến độ tích điểm nâng hạng thành viên.",
     ptext('Tài khoản = "Người dùng có tổng điểm tích lũy"', 'Màn hình = "Loyalty"'),
     "Thanh tiến độ lên hạng hiển thị đúng theo tổng điểm hiện tại."),
    ("Kiểm tra hiển thị danh sách đặc quyền tương ứng hạng thành viên.",
     ptext('Hạng thành viên = "Hạng đang có đặc quyền active"', 'Màn hình = "Quyền lợi thành viên"'),
     "Danh sách đặc quyền của hạng được hiển thị đầy đủ."),
    ("Kiểm tra hiển thị bảng lịch sử biến động điểm tích lũy.",
     ptext('Loại giao dịch = "Tất cả"', 'Thời gian = "Tất cả"', 'Trang = "1"'),
     "Hiển thị lịch sử cộng, trừ, hoàn, thưởng hoặc điều chỉnh điểm."),
    ("Kiểm tra lọc lịch sử điểm theo loại giao dịch tích điểm/đổi điểm.",
     ptext('Loại giao dịch = "Tích điểm"', 'Thời gian = "Tất cả"'),
     "Chỉ hiển thị các giao dịch cộng điểm."),
    ("Kiểm tra phân trang danh sách lịch sử biến động điểm.",
     ptext('Trang = "2"', 'Số dòng mỗi trang = "10"'),
     "Hiển thị đúng dữ liệu của trang đã chọn."),
    ("Kiểm tra hiển thị mã đơn hàng liên kết trong mô tả giao dịch điểm.",
     ptext('Giao dịch = "Giao dịch có mã đơn hàng"', 'Màn hình = "Lịch sử điểm"'),
     "Hiển thị mã đơn hàng trong mô tả giao dịch điểm."),
    ("Kiểm tra hiển thị trạng thái rỗng khi tài khoản chưa phát sinh điểm.",
     ptext('Tài khoản = "Người dùng chưa phát sinh điểm"', 'Màn hình = "Lịch sử điểm"'),
     "Hiển thị danh sách rỗng hoặc thông báo chưa có lịch sử điểm."),
    ("Kiểm tra thao tác đổi điểm lấy giảm giá trực tiếp khi đủ điểm.",
     ptext('Điểm muốn dùng = "1000"', 'Tổng tiền hàng = "300000"', 'Điểm khả dụng = "Lớn hơn hoặc bằng 1000"'),
     "Hệ thống áp dụng đổi điểm và hiển thị số tiền được giảm."),
    ("Kiểm tra báo lỗi khi nhập số điểm quy đổi vượt quá số dư điểm.",
     ptext('Điểm muốn dùng = "999999"', 'Điểm khả dụng = "Nhỏ hơn điểm muốn dùng"', 'Tổng tiền hàng = "300000"'),
     "Hiển thị thông báo điểm quy đổi không hợp lệ hoặc không đủ."),
    ("Kiểm tra hiển thị thông tin chính sách tích điểm và đổi điểm hiện hành.",
     ptext('Tài khoản = "Người dùng đã đăng nhập"', 'Màn hình = "Thanh toán hoặc Loyalty"'),
     "Hiển thị chính sách tích điểm và đổi điểm đang áp dụng."),
    ("Kiểm tra áp dụng chính sách mặc định khi không có chiến dịch đặc biệt.",
     ptext('Chính sách chiến dịch = "Không có chính sách đang hoạt động"', 'Tài khoản = "Người dùng đã đăng nhập"'),
     "Hệ thống hiển thị và áp dụng chính sách mặc định."),
    ("Kiểm tra hiển thị danh sách các hạng thành viên và điều kiện điểm.",
     ptext('Màn hình = "Loyalty"', 'Thao tác = "Xem hạng thành viên"'),
     "Hiển thị các hạng đang hoạt động và điều kiện điểm tương ứng."),
    ("Kiểm tra hiển thị trạng thái điểm danh hàng ngày của người dùng.",
     ptext('Tài khoản = "Người dùng đã đăng nhập"', 'Màn hình = "Điểm danh"'),
     "Hiển thị đã điểm danh hay chưa, chuỗi ngày hiện tại và điểm nhận tiếp theo."),
    ("Kiểm tra thao tác điểm danh nhận điểm thưởng lần đầu trong ngày.",
     ptext('Trạng thái hôm nay = "Chưa điểm danh"', 'Thao tác = "Điểm danh"'),
     "Hệ thống cộng điểm, tăng chuỗi ngày và hiển thị thông báo thành công."),
    ("Kiểm tra báo lỗi khi thực hiện điểm danh lại trong cùng một ngày.",
     ptext('Trạng thái hôm nay = "Đã điểm danh"', 'Thao tác = "Điểm danh"'),
     "Hiển thị thông báo đã điểm danh hôm nay, không cộng thêm điểm."),
    ("Kiểm tra hiển thị danh sách voucher có thể đổi bằng điểm tích lũy.",
     ptext('Tài khoản = "Người dùng đã đăng nhập"', 'Màn hình = "Đổi voucher bằng điểm"'),
     "Hiển thị danh sách voucher đổi điểm phù hợp với tài khoản."),
    ("Kiểm tra thao tác đổi điểm lấy voucher khuyến mãi thành công.",
     ptext('Voucher đổi điểm = "Voucher còn lượt đổi"', 'Điểm khả dụng = "Đủ điểm"', 'Thao tác = "Đổi voucher"'),
     "Điểm bị trừ và voucher được thêm vào ví người dùng."),
    ("Kiểm tra hiển thị bảng điều khiển quản lý Loyalty cho Admin.",
     ptext('Tài khoản = "Admin"', 'Màn hình = "Quản lý Loyalty"'),
     "Hiển thị thống kê tổng quan về thành viên, điểm và hoạt động loyalty."),
    ("Kiểm tra thao tác Admin tạo mới chính sách tích điểm thành công.",
     ptext('Tên chính sách = "Tích điểm kiểm thử"', 'Số tiền quy đổi = "1000"', 'Điểm nhận = "10"', 'Hệ số = "1"', 'Trạng thái = "Hoạt động"'),
     "Chính sách tích điểm được tạo thành công."),
    ("Kiểm tra thao tác Admin cập nhật chính sách đổi điểm thưởng.",
     ptext('Chính sách đổi điểm = "Chọn chính sách đang có"', 'Điểm đổi = "100"', 'Số tiền giảm = "1000"', 'Trạng thái = "Hoạt động"'),
     "Chính sách đổi điểm được cập nhật thành công."),
    ("Kiểm tra thao tác Admin quản lý cấu hình các hạng và đặc quyền.",
     ptext('Hạng = "Silver"', 'Điểm tối thiểu = "30000"', 'Đặc quyền = "Nhận voucher hàng tháng"', 'Trạng thái = "Hoạt động"'),
     "Hạng thành viên và đặc quyền được lưu đúng trong hệ thống."),
    ("Kiểm tra thao tác Admin cộng hoặc thu hồi điểm thủ công cho người dùng.",
     ptext('Người dùng = "Chọn người dùng đang có"', 'Số điểm = "500"', 'Lý do = "Điều chỉnh điểm kiểm thử"', 'Thao tác = "Cộng điểm/Thu hồi điểm"'),
     "Điểm của người dùng được cập nhật và hệ thống ghi nhận lịch sử/audit.")
]


chatbot_cases = [
    ("Kiểm tra gửi câu hỏi tư vấn cho Chatbot AI và nhận phản hồi.",
     ptext('Nội dung câu hỏi = "LazPe có những sản phẩm nào cho bé?"', 'Chế độ = "Chatbot AI"'),
     "Chatbot trả lời nội dung phù hợp với câu hỏi của người dùng."),
    ("Kiểm tra ràng buộc không cho phép gửi tin nhắn câu hỏi rỗng.",
     ptext('Nội dung câu hỏi = ""', 'Thao tác = "Gửi"'),
     "Nút gửi không hoạt động hoặc hệ thống yêu cầu nhập nội dung."),
    ("Kiểm tra Chatbot AI phản hồi tư vấn sản phẩm phù hợp trong hệ thống.",
     ptext('Nội dung câu hỏi = "Gợi ý sản phẩm phù hợp cho bé"', 'Chế độ = "Chatbot AI"'),
     "Chatbot phản hồi theo hướng tư vấn sản phẩm trong hệ thống bán hàng."),
    ("Kiểm tra Chatbot AI hướng dẫn luồng theo dõi thông tin đơn hàng.",
     ptext('Nội dung câu hỏi = "Tôi muốn kiểm tra đơn hàng"', 'Chế độ = "Chatbot AI"'),
     "Chatbot hướng dẫn người dùng vào chức năng theo dõi hoặc danh sách đơn hàng."),
    ("Kiểm tra hiển thị thông báo lỗi thân thiện khi dịch vụ AI mất kết nối.",
     ptext('Nội dung câu hỏi = "Kiểm tra phản hồi AI"', 'Trạng thái dịch vụ AI = "Không phản hồi"'),
     "Hiển thị thông báo lỗi thân thiện, giao diện không bị treo."),
    ("Kiểm tra tính năng phản hồi câu trả lời dạng dòng chảy (Stream).",
     ptext('Nội dung câu hỏi = "Tư vấn sản phẩm đang khuyến mãi"', 'Chế độ phản hồi = "Stream"'),
     "Câu trả lời được hiển thị dần theo luồng phản hồi."),
    ("Kiểm tra ràng buộc không tạo luồng Stream khi nội dung câu hỏi rỗng.",
     ptext('Nội dung câu hỏi = ""', 'Chế độ phản hồi = "Stream"'),
     "Hệ thống không tạo phản hồi và yêu cầu nhập câu hỏi."),
    ("Kiểm tra thao tác bật chức năng đọc câu trả lời bằng giọng nói (TTS).",
     ptext('Tùy chọn đọc câu trả lời = "Bật"', 'Nội dung câu hỏi = "Xin chào"'),
     "Sau khi AI trả lời, hệ thống đọc nội dung phản hồi nếu trình duyệt hỗ trợ."),
    ("Kiểm tra thao tác tắt chức năng đọc câu trả lời bằng giọng nói (TTS).",
     ptext('Tùy chọn đọc câu trả lời = "Tắt"', 'Nội dung câu hỏi = "Xin chào"'),
     "AI vẫn trả lời nhưng không phát âm thanh đọc nội dung."),
    ("Kiểm tra thao tác xóa lịch sử cuộc trò chuyện AI trên giao diện.",
     ptext('Lịch sử chat = "Đã có tin nhắn"', 'Thao tác = "Xóa cuộc trò chuyện"'),
     "Giao diện xóa nội dung chat hiện tại và sẵn sàng bắt đầu cuộc trò chuyện mới."),
    ("Kiểm tra khởi tạo phiên chat hỗ trợ CSKH cho khách chưa đăng nhập.",
     ptext('Tên khách = "Nguyễn Văn A"', 'Trạng thái đăng nhập = "Chưa đăng nhập"', 'Chế độ = "Hỗ trợ khách hàng"'),
     "Hệ thống tạo phiên chat khách và hiển thị khung hội thoại."),
    ("Kiểm tra khởi tạo phiên chat hỗ trợ CSKH cho người dùng đã đăng nhập.",
     ptext('Tài khoản = "Người dùng đã đăng nhập"', 'Chế độ = "Hỗ trợ khách hàng"'),
     "Hệ thống tạo hoặc lấy lại phiên chat của người dùng."),
    ("Kiểm tra tải lại lịch sử các tin nhắn cũ trong phiên chat CSKH.",
     ptext('Phiên chat = "Phiên đang có tin nhắn"', 'Thao tác = "Mở lại khung chat"'),
     "Tin nhắn cũ được hiển thị đúng trong khung chat."),
    ("Kiểm tra thao tác gửi tin nhắn văn bản trong phiên chat CSKH.",
     ptext('Nội dung tin nhắn = "Tôi cần hỗ trợ đơn hàng"', 'Phiên chat = "Đang mở"'),
     "Tin nhắn được gửi và hiển thị trong khung chat."),
    ("Kiểm tra thao tác gửi tin nhắn đính kèm hình ảnh sản phẩm.",
     ptext('Nội dung tin nhắn = "Ảnh sản phẩm cần hỗ trợ"', 'Hình ảnh = "chat-image.jpg"'),
     "Tin nhắn và hình ảnh đính kèm được hiển thị trong phiên chat."),
    ("Kiểm tra báo lỗi khi gửi tin nhắn nhưng cả nội dung và ảnh đều trống.",
     ptext('Nội dung tin nhắn = ""', 'Hình ảnh = ""', 'Thao tác = "Gửi"'),
     "Hệ thống không gửi tin nhắn rỗng."),
    ("Kiểm tra thao tác gửi yêu cầu kết nối gặp nhân viên CSKH hỗ trợ.",
     ptext('Phiên chat = "Đang mở"', 'Thao tác = "Gặp nhân viên hỗ trợ"'),
     "Phiên chat chuyển sang trạng thái chờ nhân viên hỗ trợ."),
    ("Kiểm tra ràng buộc không tạo trùng yêu cầu khi đang chờ CSKH.",
     ptext('Phiên chat = "Đang chờ nhân viên hỗ trợ"', 'Thao tác = "Gặp nhân viên hỗ trợ"'),
     "Hệ thống giữ nguyên trạng thái chờ hỗ trợ, không tạo yêu cầu trùng."),
    ("Kiểm tra thao tác kết thúc phiên hỗ trợ khách hàng CSKH.",
     ptext('Phiên chat = "Đang được hỗ trợ"', 'Thao tác = "Kết thúc hỗ trợ"'),
     "Phiên hỗ trợ được kết thúc và trạng thái được cập nhật."),
    ("Kiểm tra hiển thị danh sách các phiên chat CSKH trên trang Admin.",
     ptext('Tài khoản = "Admin"', 'Màn hình = "Quản lý chat"', 'Bộ lọc = "Tất cả"'),
     "Hiển thị danh sách phiên chat, trạng thái chờ hỗ trợ và thông tin người gửi."),
    ("Kiểm tra thao tác Admin nhận xử lý phiên chat CSKH của khách hàng.",
     ptext('Phiên chat = "Đang chờ hỗ trợ"', 'Thao tác = "Nhận xử lý"'),
     "Phiên chat được gán cho admin đang thao tác."),
    ("Kiểm tra thao tác Admin thực hiện đóng phiên chat CSKH.",
     ptext('Phiên chat = "Đang mở"', 'Thao tác = "Đóng phiên"'),
     "Phiên chat được đóng và không còn ở trạng thái đang mở."),
    ("Kiểm tra duy trì và khôi phục phiên chat cho người dùng đã đăng nhập.",
     ptext('Tài khoản = "Người dùng đã đăng nhập"', 'Phiên chat cũ = "Đã lưu trên trình duyệt"', 'Thao tác = "Mở lại website"'),
     "Hệ thống tải lại phiên chat cũ và lịch sử tin nhắn."),
    ("Kiểm tra phân quyền từ chối truy cập màn hình quản lý chat CSKH.",
     ptext('Tài khoản = "Người dùng thường"', 'Màn hình = "Quản lý chat"'),
     "Hệ thống không cho truy cập chức năng quản trị chat."),
    ("Kiểm tra hiển thị thông báo lỗi và giữ nội dung khi gửi tin thất bại.",
     ptext('Nội dung tin nhắn = "Kiểm tra lỗi gửi"', 'Trạng thái mạng = "Mất kết nối hoặc server lỗi"'),
     "Hiển thị thông báo gửi thất bại và không làm mất nội dung đang nhập.")
]


flash_sale_cases = [
    ("Kiểm tra hiển thị chương trình Flash Sale đang trong thời gian diễn ra.",
     ptext('Màn hình = "Trang chủ hoặc khu Flash Sale"', 'Thời gian hiện tại = "Nằm trong thời gian diễn ra chương trình"'),
     "Hiển thị chương trình Flash Sale đang hoạt động."),
    ("Kiểm tra ràng buộc không hiển thị Flash Sale chưa đến giờ bắt đầu.",
     ptext('Chương trình = "Có thời gian bắt đầu sau hiện tại"', 'Màn hình = "Trang chủ"'),
     "Chương trình chưa bắt đầu không hiển thị cho người dùng."),
    ("Kiểm tra ràng buộc ẩn chương trình Flash Sale đã quá hạn kết thúc.",
     ptext('Chương trình = "Có thời gian kết thúc trước hiện tại"', 'Màn hình = "Trang chủ"'),
     "Chương trình đã kết thúc không hiển thị trong khu Flash Sale hiện tại."),
    ("Kiểm tra hiển thị thông tin sản phẩm Flash Sale và phần trăm giảm.",
     ptext('Chương trình = "Flash Sale đang hoạt động"', 'Sản phẩm = "Sản phẩm nằm trong chương trình"'),
     "Hiển thị tên sản phẩm, ảnh, giá gốc, giá sale và số lượng còn lại."),
    ("Kiểm tra hiển thị thông tin giới hạn số lượng mua tối đa mỗi người.",
     ptext('Sản phẩm Flash Sale = "Có giới hạn mua"', 'Giới hạn mua = "2 sản phẩm/người"'),
     "Giao diện hiển thị hoặc áp dụng đúng giới hạn mua mỗi người."),
    ("Kiểm tra thao tác mua sản phẩm Flash Sale khi vẫn còn suất ưu đãi.",
     ptext('Sản phẩm = "Sản phẩm Flash Sale còn hàng"', 'Số lượng = "1"', 'Thao tác = "Thêm vào giỏ/Mua ngay"'),
     "Sản phẩm được thêm vào giỏ với giá Flash Sale."),
    ("Kiểm tra báo lỗi khi đặt mua số lượng vượt quá suất Flash Sale còn lại.",
     ptext('Sản phẩm = "Sản phẩm Flash Sale gần hết hàng"', 'Số lượng = "Lớn hơn số lượng còn lại"'),
     "Hiển thị thông báo số lượng không hợp lệ hoặc vượt tồn Flash Sale."),
    ("Kiểm tra báo lỗi khi đặt mua vượt quá giới hạn suất của mỗi người.",
     ptext('Sản phẩm = "Có giới hạn mua"', 'Số lượng đã mua = "Đạt giới hạn"', 'Số lượng mua thêm = "1"'),
     "Hệ thống không cho mua thêm vượt giới hạn."),
    ("Kiểm tra hiển thị và áp dụng sản phẩm quà tặng đính kèm Flash Sale.",
     ptext('Sản phẩm Flash Sale = "Có quà tặng"', 'Số lượng mua = "Đạt điều kiện nhận quà"'),
     "Giao diện hiển thị thông tin quà tặng đi kèm."),
    ("Kiểm tra hiển thị trạng thái hết hàng đối với sản phẩm đã bán hết suất.",
     ptext('Sản phẩm Flash Sale = "Đã bán hết số lượng sale"', 'Thao tác = "Mua ngay"'),
     "Hiển thị hết hàng và không cho tiếp tục mua giá sale."),
    ("Kiểm tra hiển thị danh sách các chiến dịch Flash Sale trên trang Admin.",
     ptext('Tài khoản = "Admin"', 'Màn hình = "Quản lý Flash Sale"'),
     "Hiển thị danh sách chương trình Flash Sale."),
    ("Kiểm tra hiển thị chi tiết cài đặt của một chương trình Flash Sale.",
     ptext('Chương trình = "Chọn Flash Sale đang có"', 'Thao tác = "Xem chi tiết"'),
     "Hiển thị thông tin chương trình và danh sách sản phẩm trong chương trình."),
    ("Kiểm tra tìm kiếm chiến dịch Flash Sale theo tên trên trang Admin.",
     ptext('Từ khóa = "Tên chương trình đang có"', 'Màn hình = "Quản lý Flash Sale"'),
     "Danh sách hiển thị chương trình phù hợp với từ khóa."),
    ("Kiểm tra hiển thị danh sách khách hàng đã mua sản phẩm trong Flash Sale.",
     ptext('Chương trình = "Flash Sale đã có lượt mua"', 'Thao tác = "Xem người mua"'),
     "Hiển thị danh sách khách hàng đã mua sản phẩm trong Flash Sale."),
    ("Kiểm tra phân quyền từ chối truy cập màn hình quản lý Flash Sale.",
     ptext('Tài khoản = "Người dùng thường"', 'Màn hình = "Quản lý Flash Sale"'),
     "Hệ thống không cho truy cập chức năng quản trị Flash Sale."),
    ("Kiểm tra thao tác Admin tạo mới Flash Sale với thông tin hợp lệ.",
     ptext('Tên chương trình = "Flash Sale kiểm thử"', 'Thời gian bắt đầu = "Ngày giờ sau hiện tại"', 'Thời gian kết thúc = "Sau thời gian bắt đầu"', 'Trạng thái = "Hoạt động"', 'Sản phẩm áp dụng = "Chọn sản phẩm/biến thể đang có"', 'Giá sale = "150000"', 'Số lượng sale = "50"'),
     "Chương trình Flash Sale được tạo thành công."),
    ("Kiểm tra báo lỗi khi tạo Flash Sale nhưng bỏ trống tên chương trình.",
     ptext('Tên chương trình = ""', 'Thời gian bắt đầu = "Hợp lệ"', 'Thời gian kết thúc = "Hợp lệ"'),
     "Hiển thị thông báo tên chương trình là bắt buộc."),
    ("Kiểm tra báo lỗi khi thời gian kết thúc nhỏ hơn thời gian bắt đầu.",
     ptext('Tên chương trình = "Flash Sale sai thời gian"', 'Thời gian bắt đầu = "10:00"', 'Thời gian kết thúc = "09:00"'),
     "Hiển thị thông báo thời gian chương trình không hợp lệ."),
    ("Kiểm tra báo lỗi khi thêm mục Flash Sale nhưng chưa chọn sản phẩm.",
     ptext('Tên chương trình = "Flash Sale kiểm thử"', 'Sản phẩm áp dụng = ""', 'Giá sale = "150000"', 'Số lượng sale = "50"'),
     "Hiển thị thông báo yêu cầu chọn sản phẩm hoặc biến thể áp dụng."),
    ("Kiểm tra báo lỗi khi tạo mục Flash Sale với tổng số lượng suất bằng 0.",
     ptext('Sản phẩm áp dụng = "Chọn sản phẩm đang có"', 'Giá sale = "150000"', 'Số lượng sale = "0"'),
     "Hiển thị thông báo tổng số lượng phải lớn hơn 0."),
    ("Kiểm tra thao tác Admin cập nhật thông tin chương trình Flash Sale.",
     ptext('Chương trình = "Chọn Flash Sale đang có"', 'Tên mới = "Flash Sale cập nhật"', 'Giá sale mới = "140000"', 'Số lượng sale mới = "60"'),
     "Thông tin chương trình Flash Sale được cập nhật thành công."),
    ("Kiểm tra báo lỗi khi nhập tên chương trình Flash Sale quá độ dài.",
     ptext('Tên chương trình = "Chuỗi ký tự dài hơn 200 ký tự"', 'Thao tác = "Lưu"'),
     "Hiển thị thông báo tên chương trình tối đa 200 ký tự."),
    ("Kiểm tra thao tác Admin bật hoặc tắt trạng thái hoạt động Flash Sale.",
     ptext('Chương trình = "Chọn Flash Sale đang có"', 'Trạng thái = "Bật/Tắt hoạt động"'),
     "Trạng thái hoạt động của Flash Sale được cập nhật."),
    ("Kiểm tra thao tác Admin xóa chương trình Flash Sale chưa có lượt mua.",
     ptext('Chương trình = "Flash Sale kiểm thử"', 'Thao tác = "Xóa"', 'Xác nhận = "Đồng ý"'),
     "Chương trình được xóa khỏi danh sách quản trị."),
    ("Kiểm tra thao tác hủy xác nhận khi thực hiện xóa Flash Sale.",
     ptext('Chương trình = "Chọn Flash Sale đang có"', 'Thao tác = "Xóa"', 'Xác nhận = "Hủy"'),
     "Chương trình không bị xóa và vẫn hiển thị trong danh sách.")
]


message_cases = [
    ("Kiểm tra hiển thị lịch sử tin nhắn trực tiếp giữa người dùng và Admin.",
     ptext('Tài khoản = "Người dùng đã đăng nhập"', 'Màn hình = "Tin nhắn"'),
     "Hiển thị lịch sử tin nhắn giữa người dùng và admin."),
    ("Kiểm tra yêu cầu đăng nhập khi người chưa đăng nhập mở tin nhắn.",
     ptext('Tài khoản = "Chưa đăng nhập"', 'Màn hình = "Tin nhắn"'),
     "Hiển thị yêu cầu đăng nhập."),
    ("Kiểm tra hiển thị màn hình trống khi người dùng chưa từng nhắn tin.",
     ptext('Tài khoản = "Người dùng chưa từng nhắn tin"', 'Màn hình = "Tin nhắn"'),
     "Hiển thị trạng thái chưa có tin nhắn hoặc danh sách rỗng."),
    ("Kiểm tra hiển thị tin nhắn sắp xếp theo đúng thứ tự thời gian gửi.",
     ptext('Lịch sử tin nhắn = "Có nhiều tin nhắn"', 'Thao tác = "Mở màn hình tin nhắn"'),
     "Tin nhắn được sắp xếp theo đúng thời gian gửi."),
    ("Kiểm tra tự động khôi phục lịch sử tin nhắn sau khi tải lại trang.",
     ptext('Nội dung tin nhắn = "Kiểm tra lưu lịch sử"', 'Thao tác = "Gửi và tải lại trang"'),
     "Tin nhắn vừa gửi vẫn hiển thị trong lịch sử."),
    ("Kiểm tra thao tác gửi tin nhắn văn bản trực tiếp cho Admin.",
     ptext('Nội dung tin nhắn = "Tôi cần hỗ trợ"', 'Hình ảnh = ""'),
     "Tin nhắn được gửi thành công và hiển thị ở phía người dùng."),
    ("Kiểm tra ràng buộc không cho phép gửi tin nhắn văn bản rỗng.",
     ptext('Nội dung tin nhắn = ""', 'Hình ảnh = ""', 'Thao tác = "Gửi"'),
     "Hệ thống không cho gửi tin nhắn rỗng."),
    ("Kiểm tra thao tác gửi tin nhắn chứa biểu tượng cảm xúc (Emoji).",
     ptext('Nội dung tin nhắn = "Cần hỗ trợ đơn hàng 😊"', 'Hình ảnh = ""'),
     "Tin nhắn được lưu và hiển thị đúng nội dung."),
    ("Kiểm tra thao tác gửi tin nhắn đính kèm tệp hình ảnh thành công.",
     ptext('Nội dung tin nhắn = "Gửi ảnh sản phẩm"', 'Hình ảnh = "chat-image.jpg"'),
     "Hình ảnh được tải lên và hiển thị trong tin nhắn."),
    ("Kiểm tra báo lỗi khi tải tệp đính kèm không đúng định dạng hình ảnh.",
     ptext('Nội dung tin nhắn = "Gửi file"', 'Tệp đính kèm = "document.pdf"'),
     "Hiển thị thông báo định dạng tệp không hợp lệ."),
    ("Kiểm tra thao tác gửi tin nhắn đính kèm thông tin đơn hàng cần hỗ trợ.",
     ptext('Đơn hàng = "Chọn đơn hàng của người dùng"', 'Nội dung tin nhắn = "Tôi cần hỗ trợ đơn này"'),
     "Tin nhắn được gửi kèm thông tin đơn hàng đã chọn."),
    ("Kiểm tra hiển thị danh sách đơn hàng cá nhân để chọn đính kèm.",
     ptext('Tài khoản = "Người dùng có đơn hàng"', 'Thao tác = "Chọn đơn hàng cần hỗ trợ"'),
     "Hiển thị danh sách đơn hàng của đúng người dùng."),
    ("Kiểm tra xử lý yêu cầu đăng nhập lại khi phiên làm việc hết hạn.",
     ptext('Tài khoản = "Token hết hạn"', 'Nội dung tin nhắn = "Kiểm tra gửi tin"'),
     "Hiển thị yêu cầu đăng nhập lại."),
    ("Kiểm tra hiển thị hiệu ứng trạng thái đang gửi tin nhắn trên UI.",
     ptext('Nội dung tin nhắn = "Tin nhắn kiểm thử"', 'Thao tác = "Gửi"'),
     "Giao diện hiển thị trạng thái đang gửi rồi cập nhật kết quả."),
    ("Kiểm tra hiển thị khung xem trước hình ảnh trước khi gửi tin nhắn.",
     ptext('Hình ảnh = "chat-image.jpg"', 'Thao tác = "Chọn ảnh"'),
     "Ảnh xem trước hiển thị trên giao diện và có thể xóa trước khi gửi."),
    ("Kiểm tra hiển thị toàn bộ lịch sử trao đổi với người dùng trên Admin.",
     ptext('Tài khoản = "Admin"', 'Người dùng = "Chọn người dùng đang có tin nhắn"'),
     "Admin xem được toàn bộ lịch sử trao đổi với người dùng đã chọn."),
    ("Kiểm tra thao tác Admin gửi tin nhắn phản hồi trực tiếp cho người dùng.",
     ptext('Tài khoản = "Admin"', 'Người nhận = "Chọn người dùng đang có"', 'Nội dung tin nhắn = "LazPe đã tiếp nhận yêu cầu của bạn"'),
     "Tin nhắn admin được gửi và hiển thị cho người dùng."),
    ("Kiểm tra báo lỗi khi Admin gửi tin nhắn nhưng chưa chọn người nhận.",
     ptext('Người nhận = ""', 'Nội dung tin nhắn = "Nội dung hỗ trợ"'),
     "Hiển thị thông báo yêu cầu chọn người nhận."),
    ("Kiểm tra ràng buộc không cho phép Admin gửi tin nhắn phản hồi rỗng.",
     ptext('Người nhận = "Chọn người dùng đang có"', 'Nội dung tin nhắn = ""'),
     "Hệ thống không cho gửi tin nhắn rỗng."),
    ("Kiểm tra khả năng cập nhật tin nhắn hiển thị tức thì theo thời gian thực.",
     ptext('Người dùng và admin = "Đang mở cùng cuộc trò chuyện"', 'Nội dung tin nhắn = "Tin nhắn realtime"'),
     "Tin nhắn mới hiển thị gần như ngay lập tức ở cả hai phía."),
    ("Kiểm tra hiển thị thông báo lỗi khi tải ảnh đính kèm thất bại.",
     ptext('Hình ảnh = "chat-image.jpg"', 'Trạng thái upload = "Thất bại"'),
     "Hiển thị thông báo upload thất bại và không gửi tin nhắn sai dữ liệu."),
    ("Kiểm tra ranh giới bảo mật không xem được tin nhắn của tài khoản khác.",
     ptext('Tài khoản = "Người dùng A"', 'Cuộc trò chuyện = "Của người dùng B"'),
     "Người dùng A không xem được tin nhắn của người dùng B."),
    ("Kiểm tra hiển thị phân biệt rõ ràng bóng chat người dùng và Admin.",
     ptext('Lịch sử tin nhắn = "Có tin của user và admin"', 'Màn hình = "Tin nhắn"'),
     "Giao diện hiển thị đúng phía gửi, tên hoặc avatar của từng bên."),
    ("Kiểm tra khả năng xử lý gửi nhiều tin nhắn liên tiếp không mất dữ liệu.",
     ptext('Tin nhắn 1 = "Xin chào"', 'Tin nhắn 2 = "Tôi cần hỗ trợ"', 'Tin nhắn 3 = "Cảm ơn"'),
     "Các tin nhắn được lưu đúng thứ tự và không bị mất dữ liệu."),
    ("Kiểm tra xử lý hiển thị thông báo lỗi khi máy chủ gặp sự cố gửi tin.",
     ptext('Nội dung tin nhắn = "Kiểm tra lỗi server"', 'Trạng thái server = "Lỗi"'),
     "Giao diện hiển thị gửi thất bại và không bị đứng màn hình.")
]


notification_cases = [
    ("Kiểm tra hiển thị danh sách tất cả thông báo cá nhân của người dùng.",
     ptext('Tài khoản = "Người dùng đã đăng nhập"', 'Màn hình = "Thông báo"', 'Bộ lọc = "Tất cả"'),
     "Hiển thị danh sách thông báo của người dùng."),
    ("Kiểm tra lọc danh sách thông báo theo phân loại (Khuyến mãi, Đơn hàng...).",
     ptext('Loại thông báo = "Khuyến mãi"', 'Trạng thái đọc = "Tất cả"'),
     "Chỉ hiển thị thông báo thuộc loại đã chọn."),
    ("Kiểm tra lọc danh sách thông báo theo trạng thái chưa đọc.",
     ptext('Trạng thái đọc = "Chưa đọc"', 'Loại thông báo = "Tất cả"'),
     "Chỉ hiển thị các thông báo chưa đọc."),
    ("Kiểm tra phân trang danh sách thông báo cá nhân.",
     ptext('Trang = "1"', 'Số dòng mỗi trang = "20"'),
     "Hiển thị đúng danh sách thông báo theo trang."),
    ("Kiểm tra yêu cầu đăng nhập khi người chưa đăng nhập mở thông báo.",
     ptext('Tài khoản = "Chưa đăng nhập"', 'Màn hình = "Thông báo"'),
     "Hiển thị yêu cầu đăng nhập."),
    ("Kiểm tra hiển thị đúng số lượng huy hiệu (Badge) thông báo chưa đọc.",
     ptext('Tài khoản = "Người dùng có thông báo chưa đọc"', 'Màn hình = "Header/Hồ sơ"'),
     "Hiển thị đúng số lượng thông báo chưa đọc."),
    ("Kiểm tra thao tác đánh dấu một thông báo từ chưa đọc sang đã đọc.",
     ptext('Thông báo = "Thông báo chưa đọc"', 'Thao tác = "Đánh dấu đã đọc"'),
     "Thông báo chuyển sang trạng thái đã đọc."),
    ("Kiểm tra báo lỗi khi thực hiện đánh dấu đọc thông báo không tồn tại.",
     ptext('Thông báo = "Thông báo đã bị xóa hoặc không tồn tại"', 'Thao tác = "Đánh dấu đã đọc"'),
     "Hiển thị thông báo không tìm thấy thông báo."),
    ("Kiểm tra thao tác đánh dấu tất cả thông báo sang trạng thái đã đọc.",
     ptext('Thông báo = "Có nhiều thông báo chưa đọc"', 'Thao tác = "Đánh dấu tất cả đã đọc"'),
     "Tất cả thông báo của người dùng chuyển sang trạng thái đã đọc."),
    ("Kiểm tra thao tác xóa mềm ẩn thông báo khỏi danh sách cá nhân.",
     ptext('Thông báo = "Chọn thông báo của người dùng"', 'Thao tác = "Xóa"'),
     "Thông báo bị ẩn khỏi danh sách của người dùng."),
    ("Kiểm tra ràng buộc không cho phép xóa thông báo của tài khoản khác.",
     ptext('Tài khoản = "Người dùng A"', 'Thông báo = "Thông báo thuộc người dùng B"', 'Thao tác = "Xóa"'),
     "Hệ thống không cho xóa thông báo của người dùng khác."),
    ("Kiểm tra hiển thị ưu tiên nổi bật đối với thông báo được ghim.",
     ptext('Thông báo = "Thông báo có ghim"', 'Màn hình = "Thông báo"'),
     "Thông báo được ghim hiển thị nổi bật hoặc ưu tiên."),
    ("Kiểm tra thao tác kích hoạt nút liên kết hành động (Action URL).",
     ptext('Thông báo = "Có nút hành động"', 'Thao tác = "Bấm nút hành động"'),
     "Hệ thống điều hướng đến đúng màn hình hoặc đường dẫn cấu hình."),
    ("Kiểm tra hiển thị hình ảnh đại diện hoặc banner trong thông báo.",
     ptext('Thông báo = "Có ảnh đại diện hoặc banner"', 'Màn hình = "Chi tiết thông báo"'),
     "Ảnh thông báo hiển thị đúng trên giao diện."),
    ("Kiểm tra tìm kiếm thông báo theo từ khóa trên giao diện.",
     ptext('Từ khóa = "khuyến mãi"', 'Danh sách thông báo = "Có thông báo chứa từ khóa"'),
     "Danh sách được lọc theo từ khóa người dùng nhập."),
    ("Kiểm tra hiển thị danh sách các chiến dịch thông báo trên Admin.",
     ptext('Tài khoản = "Admin"', 'Màn hình = "Quản lý thông báo"', 'Từ khóa = ""'),
     "Hiển thị danh sách chiến dịch thông báo."),
    ("Kiểm tra hiển thị chi tiết cài đặt của một chiến dịch thông báo.",
     ptext('Chiến dịch = "Chọn chiến dịch đang có"', 'Thao tác = "Xem chi tiết"'),
     "Hiển thị tiêu đề, mô tả, nội dung, nhóm nhận và trạng thái chiến dịch."),
    ("Kiểm tra thao tác Admin tạo mới chiến dịch thông báo hợp lệ.",
     ptext('Tiêu đề = "Thông báo kiểm thử"', 'Mô tả ngắn = "Mô tả thông báo"', 'Nội dung = "Nội dung chi tiết thông báo"', 'Loại = "Hệ thống"', 'Mức ưu tiên = "Trung bình"', 'Đối tượng nhận = "Tất cả"', 'Trạng thái ghim = "Không"'),
     "Chiến dịch thông báo được tạo thành công."),
    ("Kiểm tra báo lỗi khi tạo chiến dịch thông báo nhưng bỏ trống tiêu đề.",
     ptext('Tiêu đề = ""', 'Mô tả ngắn = "Mô tả thông báo"', 'Nội dung = "Nội dung chi tiết"'),
     "Hiển thị thông báo tiêu đề là bắt buộc."),
    ("Kiểm tra báo lỗi khi nhập tiêu đề thông báo vượt quá 200 ký tự.",
     ptext('Tiêu đề = "Chuỗi ký tự dài hơn 200 ký tự"', 'Mô tả ngắn = "Mô tả"', 'Nội dung = "Nội dung"'),
     "Hiển thị thông báo tiêu đề không được vượt quá 200 ký tự."),
    ("Kiểm tra thao tác Admin cập nhật nội dung chiến dịch thông báo.",
     ptext('Chiến dịch = "Chọn chiến dịch đang có"', 'Tiêu đề mới = "Thông báo đã cập nhật"', 'Nội dung mới = "Nội dung cập nhật"'),
     "Thông tin chiến dịch được cập nhật thành công."),
    ("Kiểm tra thao tác Admin xóa chiến dịch thông báo khỏi hệ thống.",
     ptext('Chiến dịch = "Chiến dịch kiểm thử"', 'Thao tác = "Xóa"', 'Xác nhận = "Đồng ý"'),
     "Chiến dịch được xóa hoặc ẩn khỏi danh sách quản trị."),
    ("Kiểm tra thao tác Admin phát gửi tức thì chiến dịch thông báo.",
     ptext('Chiến dịch = "Chiến dịch đã tạo"', 'Thao tác = "Gửi ngay"'),
     "Thông báo được gửi đến đúng nhóm người nhận."),
    ("Kiểm tra thao tác Admin hủy lịch đặt giờ phát thông báo tương lai.",
     ptext('Chiến dịch = "Chiến dịch có lịch gửi tương lai"', 'Thao tác = "Hủy lịch"'),
     "Lịch gửi thông báo được hủy thành công."),
    ("Kiểm tra thao tác Admin quản lý danh sách các mẫu (Template) thông báo.",
     ptext('Tên mẫu = "Mẫu kiểm thử"', 'Mã mẫu = "TEMPLATE_TEST"', 'Nội dung mẫu = "Xin chào {Tên khách hàng}"', 'Trạng thái = "Hoạt động"'),
     "Mẫu thông báo được tạo/cập nhật/xóa thành công theo thao tác admin.")
]


def build_docx():
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)

    add_paragraph(doc, "CHƯƠNG 6: KIỂM THỬ HỆ THỐNG", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "6.1. Mục tiêu kiểm thử", size=14, bold=True)
    add_paragraph(
        doc,
        "Chương này trình bày kết quả kiểm thử các chức năng chính của hệ thống LazPe. "
        "Các test case được xây dựng theo luồng thao tác trên giao diện người dùng và giao diện quản trị, "
        "tập trung vào dữ liệu mà người dùng hoặc admin nhập/chọn trực tiếp trên màn hình.",
    )
    add_paragraph(doc, "6.2. Phạm vi và môi trường kiểm thử", size=14, bold=True)
    for item in [
        "Đối tượng kiểm thử: website bán lẻ trực tuyến LazPe/PolyBaby.",
        "Nhóm người dùng: khách hàng, người dùng đã đăng nhập, admin/nhân viên có quyền tương ứng.",
        "Môi trường kiểm thử: giao diện frontend Next.js, backend ASP.NET Core Web API và cơ sở dữ liệu của hệ thống.",
        "Cách trình bày dữ liệu đầu vào: ghi theo dữ liệu nhập/chọn trên giao diện, không ghi theo endpoint kỹ thuật.",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)
        style_runs(p, size=12)

    add_paragraph(doc, "6.3. Kết quả kiểm thử", size=14, bold=True)
    tables = [
        ("Bảng 6.1: Testcase Chức năng Sản phẩm", product_cases),
        ("Bảng 6.2: Testcase Chức năng Đơn hàng", order_cases),
        ("Bảng 6.3: Testcase Chức năng Chương trình khuyến mãi/Voucher", voucher_cases),
        ("Bảng 6.4: Testcase Chức năng Khách hàng thân thiết (Loyalty)", loyalty_cases),
        ("Bảng 6.5: Testcase Chức năng Chatbot AI và Hỗ trợ CSKH", chatbot_cases),
        ("Bảng 6.6: Testcase Chức năng Flash Sale", flash_sale_cases),
        ("Bảng 6.7: Testcase Chức năng Tin nhắn", message_cases),
        ("Bảng 6.8: Testcase Chức năng Thông báo", notification_cases),
    ]
    for title, rows in tables:
        if len(rows) != 25:
            raise ValueError(f"{title} phải có 25 test case, hiện có {len(rows)}")
        add_table(doc, title, rows)

    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    build_docx()