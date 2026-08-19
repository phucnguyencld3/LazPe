from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
MD_PATH = ROOT / "docs" / "chuong-6-testcase-lazpe.md"
OUT_PATH = ROOT / "docs" / "chuong-6-testcase-lazpe.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_width(table, widths):
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)


def style_paragraph(paragraph, font_name="Times New Roman", size=12, bold=False, italic=False):
    for run in paragraph.runs:
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic


def add_text_paragraph(doc, text):
    p = doc.add_paragraph()
    if text.startswith("- "):
        p.style = "List Bullet"
        p.add_run(text[2:])
    else:
        p.add_run(text)
    style_paragraph(p, size=12)
    p.paragraph_format.space_after = Pt(4)
    return p


def split_markdown_table(lines, start):
    rows = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        line = lines[i].strip()
        parts = [part.strip().replace("\\|", "|") for part in line.strip("|").split("|")]
        rows.append(parts)
        i += 1
    return rows, i


def add_testcase_table(doc, rows):
    header = rows[0]
    body = rows[2:]
    table = doc.add_table(rows=1, cols=len(header))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    widths = [0.45, 2.45, 2.35, 3.05, 0.7]

    hdr_cells = table.rows[0].cells
    for idx, text in enumerate(header):
        hdr_cells[idx].text = text
        set_cell_shading(hdr_cells[idx], "E8EEF5")
        set_cell_margins(hdr_cells[idx])
        hdr_cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for p in hdr_cells[idx].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            style_paragraph(p, size=10.5, bold=True)
    set_repeat_table_header(table.rows[0])

    for row in body:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            cells[idx].text = text.replace("`", "")
            set_cell_margins(cells[idx])
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cells[idx].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx in (0, 4) else WD_ALIGN_PARAGRAPH.LEFT
                style_paragraph(p, size=10)
        if row and row[-1].lower() == "pass":
            for p in cells[-1].paragraphs:
                for run in p.runs:
                    run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
                    run.bold = True

    set_table_width(table, widths)
    doc.add_paragraph()


def build_docx():
    lines = MD_PATH.read_text(encoding="utf-8").splitlines()
    doc = Document()

    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)

    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue

        if line.startswith("# "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(line[2:].upper())
            style_paragraph(p, size=16, bold=True)
            p.paragraph_format.space_after = Pt(10)
        elif line.startswith("## "):
            p = doc.add_paragraph()
            p.add_run(line[3:])
            style_paragraph(p, size=14, bold=True)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(6)
        elif line.startswith("### "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(line[4:])
            style_paragraph(p, size=12, bold=True, italic=True)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(4)
        elif line.startswith("|"):
            rows, i = split_markdown_table(lines, i)
            add_testcase_table(doc, rows)
            continue
        else:
            add_text_paragraph(doc, line)
        i += 1

    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    build_docx()
