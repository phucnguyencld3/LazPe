from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from generate_erd_module_docs import MODULES, module_script, parse_erd


ROOT = Path(__file__).resolve().parents[1]
OUT_DOCX = Path(
    r"C:\Users\Thanh\OneDrive\Desktop\FPLOLI\FV Team - Tach ERD Mermaid va mo ta chi tiet LazPe.docx"
)


def set_font(run, name: str = "Calibri", size: float = 10.5, bold: bool = False, color: str | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def shade_paragraph(paragraph, fill: str = "F6F8FA") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_font(run, size=16 if level == 1 else 13 if level == 2 else 11.5, bold=True, color="2E74B5")


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    set_font(p.add_run(text), size=10.5)


def add_label_bullet(doc: Document, label: str, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    set_font(p.add_run(f"{label}: "), size=10.5, bold=True, color="1F4D78")
    set_font(p.add_run(text), size=10.5)


def add_code_block(doc: Document, code: str) -> None:
    for line in code.rstrip().splitlines():
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Inches(0.12)
        shade_paragraph(p)
        set_font(p.add_run(line), name="Consolas", size=7.2, color="1F2937")


def setup_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08

    for name, size in [("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11.5)]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string("2E74B5" if name != "Heading 3" else "1F4D78")
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(4)

    return doc


def main() -> None:
    tables, rels = parse_erd()
    doc = setup_document()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(title.add_run("TÁCH ERD VẬT LÝ THEO MÔ ĐUN"), size=18, bold=True, color="0B2545")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(
        subtitle.add_run("Script Mermaid và mô tả chi tiết cho dự án LazPe"),
        size=12.5,
        bold=True,
        color="2E74B5",
    )

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(
        note.add_run(
            "Nguồn: ERD vật lý full sinh từ ApplicationDbContext, migration snapshot, model SQL Server và MongoDB."
        ),
        size=9.5,
        color="555555",
    )

    add_heading(doc, "1. Cách sử dụng", 1)
    doc.add_paragraph(
        "Mỗi mô đun bên dưới gồm script Mermaid để copy vào draw.io hoặc Mermaid Live Editor, "
        "sau đó là phần mô tả chi tiết gồm mục đích, chức năng, quan hệ dữ liệu và giải thích nghiệp vụ."
    )
    add_bullet(doc, "Các bảng trung tâm có thể lặp lại ở nhiều mô đun để từng ERD nhỏ đọc được độc lập.")
    add_bullet(doc, "PK là khóa chính, FK là khóa ngoại, PK/FK là cột vừa thuộc khóa chính vừa là khóa ngoại.")
    add_bullet(doc, "Các collection MongoDB không enforce FK vật lý; các mã UserId/ProductId là liên kết logic.")

    add_heading(doc, "2. Danh sách ERD mô đun", 1)
    for module in MODULES:
        add_bullet(doc, f"{module['code']} - {module['title']}")

    add_heading(doc, "3. Chi tiết từng ERD mô đun", 1)
    for module in MODULES:
        add_heading(doc, f"{module['code']}. {module['title']}", 2)

        add_heading(doc, "Script Mermaid", 3)
        add_code_block(doc, module_script(module, tables, rels))

        add_heading(doc, "Mục đích", 3)
        doc.add_paragraph(str(module["purpose"]))

        add_heading(doc, "Chức năng của mô đun", 3)
        for item in module["functions"]:
            add_bullet(doc, str(item))

        add_heading(doc, "Quan hệ dữ liệu trong sơ đồ", 3)
        for item in module["relationships"]:
            add_bullet(doc, str(item))

        add_heading(doc, "Giải thích nghiệp vụ", 3)
        doc.add_paragraph(str(module["explanation"]))

    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    main()
